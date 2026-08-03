from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from ..util import atomic_write_bytes, sha256_bytes
from .query import iter_appendix_sheets

_CELL_LIMIT = 200


def _stringify(value: Any) -> str:
    if value is None:
        return ""
    text = str(value)
    if len(text) > _CELL_LIMIT:
        return text[: _CELL_LIMIT - 1] + "…"
    return text


def _table_matrix(rows: list[dict[str, Any]]) -> tuple[list[str], list[list[str]]]:
    fieldnames = sorted({key for row in rows for key in row})
    body = [[_stringify(row.get(key, "")) for key in fieldnames] for row in rows]
    return fieldnames, body


def render_docx(query: dict[str, Any]) -> bytes | None:
    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_BREAK  # type: ignore
    except ImportError:
        return None
    document = Document()
    document.add_heading("NeuroAI observatory publication", level=1)
    document.add_paragraph(
        "Generators are views: this document restates canonical JSON projections only. "
        "It does not establish scientific truth, regulatory authorization, or institutional authority."
    )
    document.add_paragraph(f"release_sha256: {query['release_sha256']}")
    document.add_paragraph(f"release_version: {query['metadata'].get('version', 'UNRESOLVED')}")
    document.add_paragraph(f"release_kind: {query.get('release_kind', 'UNRESOLVED')}")
    document.add_paragraph(f"query_depth: {query.get('depth', 'UNRESOLVED')}")

    document.add_heading("Withheld claims", level=2)
    for claim in query["withheld_claims"]:
        document.add_paragraph(claim, style="List Bullet")

    document.add_heading("Release summary", level=2)
    for row in query["rows"].get("release_summary", []):
        document.add_paragraph(
            ", ".join(f"{key}={row.get(key)}" for key in sorted(row)),
        )

    counts_key = "successor_counts" if "successor_counts" in query["rows"] else "coverage_counts"
    if counts_key in query["rows"]:
        document.add_heading("Counts", level=2)
        for row in query["rows"][counts_key]:
            document.add_paragraph(f"{row.get('metric')}: {row.get('value')}", style="List Bullet")

    appendices = iter_appendix_sheets(query)
    if appendices:
        document.add_heading("Publication appendices", level=2)
        document.add_paragraph(
            "The following tables project first-class query sheets present on the canonical release. "
            "Missing canonical sections are omitted."
        )
        for sheet_name, rows in appendices:
            document.add_heading(sheet_name.replace("_", " "), level=3)
            document.add_paragraph(f"Rows: {len(rows)}")
            fieldnames, body = _table_matrix(rows)
            table = document.add_table(rows=1 + len(body), cols=len(fieldnames))
            table.style = "Table Grid"
            header_cells = table.rows[0].cells
            for index, name in enumerate(fieldnames):
                header_cells[index].text = name
            for row_index, values in enumerate(body, start=1):
                cells = table.rows[row_index].cells
                for col_index, value in enumerate(values):
                    cells[col_index].text = value
            # Keep large multi-table documents navigable.
            if sheet_name != appendices[-1][0]:
                document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    document.add_heading("Boundary", level=2)
    document.add_paragraph(
        "Boundary: Document restates canonical JSON projections only; it does not establish substantive authority. "
        "Generators are views."
    )
    buffer_path_bytes = io.BytesIO()
    document.save(buffer_path_bytes)
    return bytes(buffer_path_bytes.getvalue())


def write_docx(query: dict[str, Any], output: Path) -> dict[str, Any]:
    payload = render_docx(query)
    if payload is None:
        text = (
            "NeuroAI publication Word stub\n"
            "Native .docx requires optional dependency python-docx (extra: products).\n"
            f"release_sha256: {query['release_sha256']}\n"
        )
        atomic_write_bytes(output, text.encode("utf-8"))
        return {
            "output": str(output),
            "sha256": sha256_bytes(text.encode("utf-8")),
            "bytes": len(text.encode("utf-8")),
            "format": "docx-stub-text",
            "boundary": "Stub records release identity until python-docx is installed.",
        }
    atomic_write_bytes(output, payload)
    return {
        "output": str(output),
        "sha256": sha256_bytes(payload),
        "bytes": len(payload),
        "format": "python-docx-native",
        "boundary": "DOCX is a deterministic projection of canonical JSON; no new findings.",
    }
