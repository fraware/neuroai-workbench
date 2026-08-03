#!/usr/bin/env python3
"""Create a loss-aware v2.2.0 migration baseline and reconciliation package.

This command treats the consolidated XLSX and DOCX as immutable migration inputs.
It exports workbook sheets to deterministic JSON, inventories report structure,
computes archive/file identities, and fails closed on baseline-count drift.

It does not update substantive claims, refresh public sources, or silently choose
between conflicting records. Those activities belong to later v2.3 phases.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import os
import re
import sys
import zipfile
from collections import Counter
from datetime import date, datetime
from decimal import Decimal
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter

BOUNDARY = (
    "Migration records structure, identity, and parity only. It does not establish "
    "scientific truth, regulatory status, clinical value, conformance, or currentness."
)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def canonical_bytes(value: Any) -> bytes:
    return json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")


def sha256_value(value: Any) -> str:
    return hashlib.sha256(canonical_bytes(value)).hexdigest()


def json_safe(value: Any) -> Any:
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    if isinstance(value, Decimal):
        return str(value)
    if isinstance(value, bytes):
        return {"encoding": "hex", "value": value.hex()}
    if isinstance(value, float):
        if value != value:
            return "NaN"
        if value == float("inf"):
            return "Infinity"
        if value == float("-inf"):
            return "-Infinity"
    return value


def clean_header(value: Any, column_index: int, used: set[str]) -> str:
    raw = str(value).strip() if value is not None else ""
    base = raw or f"column_{get_column_letter(column_index)}"
    base = re.sub(r"\s+", " ", base)
    candidate = base
    suffix = 2
    while candidate in used:
        candidate = f"{base}__{suffix}"
        suffix += 1
    used.add(candidate)
    return candidate


def load_dispositions(path: Path) -> dict[str, dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle))
    result: dict[str, dict[str, str]] = {}
    for row in rows:
        name = str(row.get("sheet", "")).strip()
        if not name:
            continue
        result[name] = {key: str(value or "") for key, value in row.items()}
    return result


def export_workbook(
    workbook_path: Path,
    output_root: Path,
    dispositions: dict[str, dict[str, str]],
) -> tuple[list[dict[str, Any]], dict[str, list[dict[str, Any]]]]:
    wb = load_workbook(workbook_path, data_only=False, read_only=False)
    exports_dir = output_root / "canonical_sheet_exports"
    exports_dir.mkdir(parents=True, exist_ok=True)
    inventory: list[dict[str, Any]] = []
    rows_by_sheet: dict[str, list[dict[str, Any]]] = {}

    for sheet_index, ws in enumerate(wb.worksheets):
        max_row = ws.max_row or 0
        max_col = ws.max_column or 0
        used: set[str] = set()
        headers = [clean_header(ws.cell(1, col).value, col, used) for col in range(1, max_col + 1)] if max_row else []
        rows: list[dict[str, Any]] = []
        formula_cells: list[dict[str, Any]] = []
        nonempty_cells = 0

        for row_index in range(2, max_row + 1):
            record: dict[str, Any] = {"_source_row": row_index}
            has_value = False
            for col_index, header in enumerate(headers, start=1):
                cell = ws.cell(row_index, col_index)
                value = json_safe(cell.value)
                if value is not None and value != "":
                    has_value = True
                    nonempty_cells += 1
                record[header] = value
                if isinstance(cell.value, str) and cell.value.startswith("="):
                    formula_cells.append({"address": cell.coordinate, "formula": cell.value})
            if has_value:
                rows.append(record)

        nonempty_cells += sum(1 for col in range(1, max_col + 1) if ws.cell(1, col).value not in (None, ""))
        export = {
            "metadata": {
                "sheet": ws.title,
                "sheet_index": sheet_index,
                "source_workbook": workbook_path.name,
                "source_workbook_sha256": sha256_file(workbook_path),
                "max_row": max_row,
                "max_column": max_col,
                "data_row_count": len(rows),
                "formula_cell_count": len(formula_cells),
                "nonempty_cell_count": nonempty_cells,
                "disposition": dispositions.get(ws.title, {}),
                "boundary": BOUNDARY,
            },
            "columns": headers,
            "rows": rows,
            "formula_cells": formula_cells,
        }
        export["metadata"]["canonical_export_sha256"] = sha256_value(export)
        target = exports_dir / f"{sheet_index:02d}_{ws.title}.json"
        target.write_text(json.dumps(export, ensure_ascii=False, indent=2), encoding="utf-8")
        rows_by_sheet[ws.title] = rows
        inventory.append(
            {
                "sheet_index": sheet_index,
                "sheet": ws.title,
                "max_row": max_row,
                "max_column": max_col,
                "data_row_count": len(rows),
                "formula_cell_count": len(formula_cells),
                "nonempty_cell_count": nonempty_cells,
                "export_path": str(target.relative_to(output_root)).replace(os.sep, "/"),
                "export_sha256": sha256_file(target),
                **dispositions.get(ws.title, {}),
            }
        )
    return inventory, rows_by_sheet


def inventory_docx(docx_path: Path, output_root: Path) -> dict[str, Any]:
    doc = Document(docx_path)
    headings: list[dict[str, Any]] = []
    paragraphs: list[dict[str, Any]] = []
    style_counts: Counter[str] = Counter()
    for index, paragraph in enumerate(doc.paragraphs):
        style = paragraph.style.name if paragraph.style is not None else "UNSTYLED"
        text = paragraph.text
        style_counts[style] += 1
        paragraphs.append({"index": index, "style": style, "text": text})
        if style.lower().startswith("heading") or style.lower() in {"title", "subtitle"}:
            headings.append({"paragraph_index": index, "style": style, "text": text})

    tables = []
    for index, table in enumerate(doc.tables):
        row_count = len(table.rows)
        col_count = max((len(row.cells) for row in table.rows), default=0)
        first_row = [cell.text for cell in table.rows[0].cells] if table.rows else []
        tables.append(
            {
                "table_index": index,
                "row_count": row_count,
                "column_count": col_count,
                "first_row": first_row,
            }
        )

    inventory = {
        "metadata": {
            "source_docx": docx_path.name,
            "source_docx_sha256": sha256_file(docx_path),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "section_count": len(doc.sections),
            "inline_shape_count": len(doc.inline_shapes),
            "heading_count": len(headings),
            "boundary": BOUNDARY,
        },
        "core_properties": {
            "title": doc.core_properties.title,
            "subject": doc.core_properties.subject,
            "author": doc.core_properties.author,
            "keywords": doc.core_properties.keywords,
            "comments": doc.core_properties.comments,
            "created": json_safe(doc.core_properties.created),
            "modified": json_safe(doc.core_properties.modified),
            "last_modified_by": doc.core_properties.last_modified_by,
            "revision": doc.core_properties.revision,
        },
        "style_counts": dict(sorted(style_counts.items())),
        "headings": headings,
        "tables": tables,
        "paragraphs": paragraphs,
    }
    inventory["metadata"]["inventory_sha256"] = sha256_value(inventory)
    path = output_root / "report_structure_inventory.json"
    path.write_text(json.dumps(inventory, ensure_ascii=False, indent=2), encoding="utf-8")
    return inventory


def inventory_archive(zip_path: Path, output_root: Path) -> dict[str, Any]:
    members: list[dict[str, Any]] = []
    with zipfile.ZipFile(zip_path) as archive:
        for info in archive.infolist():
            if info.is_dir():
                continue
            members.append(
                {
                    "path": info.filename,
                    "size_bytes": info.file_size,
                    "compressed_size_bytes": info.compress_size,
                    "crc32": f"{info.CRC:08x}",
                }
            )
    members.sort(key=lambda item: item["path"])
    inventory = {
        "metadata": {
            "archive": zip_path.name,
            "archive_sha256": sha256_file(zip_path),
            "file_count": len(members),
            "uncompressed_size_bytes": sum(item["size_bytes"] for item in members),
            "compressed_member_bytes": sum(item["compressed_size_bytes"] for item in members),
            "boundary": BOUNDARY,
        },
        "members": members,
    }
    inventory["metadata"]["inventory_sha256"] = sha256_value(inventory)
    path = output_root / "archive_inventory.json"
    path.write_text(json.dumps(inventory, ensure_ascii=False, indent=2), encoding="utf-8")
    return inventory


def normalize_int(value: Any) -> int | None:
    if value is None or value == "":
        return None
    if isinstance(value, bool):
        return int(value)
    if isinstance(value, (int, float)):
        return int(value)
    try:
        return int(str(value).strip())
    except ValueError:
        return None


def reconcile(
    contract: dict[str, Any],
    workbook_inventory: list[dict[str, Any]],
    rows_by_sheet: dict[str, list[dict[str, Any]]],
    report_inventory: dict[str, Any],
    archive_inventory: dict[str, Any],
) -> dict[str, Any]:
    checks: list[dict[str, Any]] = []

    def add(check_id: str, category: str, expected: Any, observed: Any, detail: str = "") -> None:
        checks.append(
            {
                "check_id": check_id,
                "category": category,
                "expected": expected,
                "observed": observed,
                "status": "PASS" if observed == expected else "FAIL",
                "detail": detail,
            }
        )

    add("WB-SHEET-COUNT", "WORKBOOK_STRUCTURE", contract["expected_sheet_count"], len(workbook_inventory))
    sheet_counts = {item["sheet"]: item["data_row_count"] for item in workbook_inventory}
    for sheet, expected in contract["expected_key_counts"].items():
        add(f"COUNT-{sheet}", "KEY_COUNT", expected, sheet_counts.get(sheet))

    summary_rows = rows_by_sheet.get("Assessment_Summary", [])
    summary_by_system = {str(row.get("system")): row for row in summary_rows}
    for system, expected_values in contract["assessment_summary"].items():
        row = summary_by_system.get(system, {})
        for field, expected in expected_values.items():
            observed = normalize_int(row.get(field))
            add(f"ASM-{system}-{field}", "ASSESSMENT_SUMMARY", expected, observed)

    add("DOCX-PARAGRAPH-COUNT", "REPORT_STRUCTURE", 4267, report_inventory["metadata"]["paragraph_count"])
    add("DOCX-TABLE-COUNT", "REPORT_STRUCTURE", 349, report_inventory["metadata"]["table_count"])
    add("DOCX-SECTION-COUNT", "REPORT_STRUCTURE", 107, report_inventory["metadata"]["section_count"])
    add("ARCHIVE-FILE-COUNT", "ARCHIVE_STRUCTURE", 1344, archive_inventory["metadata"]["file_count"])

    failures = [item for item in checks if item["status"] == "FAIL"]
    report = {
        "metadata": {
            "report": "V2.2.0 migration reconciliation",
            "contract_version": contract["contract_version"],
            "baseline_release": contract["baseline_release"],
            "target_release": contract["target_release"],
            "status": "PASS" if not failures else "FAIL",
            "check_count": len(checks),
            "pass_count": len(checks) - len(failures),
            "fail_count": len(failures),
            "boundary": BOUNDARY,
        },
        "checks": checks,
        "failures": failures,
        "next_phase_gate": {
            "static_import_parity_allowed": not failures,
            "substantive_refresh_allowed": False,
            "reason": (
                "Import parity checks passed. A separate controlled successor-update phase is still required."
                if not failures
                else "Baseline reconciliation failures must be resolved before canonical migration."
            ),
        },
    }
    report["metadata"]["report_sha256"] = sha256_value(report)
    return report


def write_csv(path: Path, rows: Iterable[dict[str, Any]]) -> None:
    materialized = list(rows)
    fieldnames: list[str] = []
    for row in materialized:
        for key in row:
            if key not in fieldnames:
                fieldnames.append(key)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(materialized)


def build_manifest(output_root: Path, inputs: dict[str, Path]) -> dict[str, Any]:
    files = []
    for path in sorted(output_root.rglob("*")):
        if not path.is_file() or path.name == "MIGRATION_PACKAGE_MANIFEST.json":
            continue
        files.append(
            {
                "path": str(path.relative_to(output_root)).replace(os.sep, "/"),
                "size_bytes": path.stat().st_size,
                "sha256": sha256_file(path),
            }
        )
    manifest = {
        "metadata": {
            "title": "NeuroAI v2.2.0 to v2.3.0-dev migration package",
            "file_count": len(files),
            "boundary": BOUNDARY,
        },
        "inputs": {
            key: {"path": str(value), "size_bytes": value.stat().st_size, "sha256": sha256_file(value)}
            for key, value in inputs.items()
        },
        "files": files,
    }
    manifest["metadata"]["manifest_sha256"] = sha256_value(manifest)
    return manifest


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--workbook", type=Path, required=True)
    parser.add_argument("--report", type=Path, required=True)
    parser.add_argument("--archive", type=Path, required=True)
    parser.add_argument("--sheet-map", type=Path, required=True)
    parser.add_argument("--contract", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args(argv)

    for path in (args.workbook, args.report, args.archive, args.sheet_map, args.contract):
        if not path.is_file():
            parser.error(f"Required input does not exist: {path}")
    args.output.mkdir(parents=True, exist_ok=True)

    contract = json.loads(args.contract.read_text(encoding="utf-8"))
    dispositions = load_dispositions(args.sheet_map)
    workbook_inventory, rows_by_sheet = export_workbook(args.workbook, args.output, dispositions)
    write_csv(args.output / "workbook_sheet_inventory.csv", workbook_inventory)
    (args.output / "workbook_sheet_inventory.json").write_text(
        json.dumps({"sheets": workbook_inventory, "boundary": BOUNDARY}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    report_inventory = inventory_docx(args.report, args.output)
    archive_inventory = inventory_archive(args.archive, args.output)
    reconciliation = reconcile(contract, workbook_inventory, rows_by_sheet, report_inventory, archive_inventory)
    (args.output / "MIGRATION_RECONCILIATION_REPORT_v2.3.0-rc1.json").write_text(
        json.dumps(reconciliation, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    write_csv(args.output / "MIGRATION_RECONCILIATION_CHECKS_v2.3.0-rc1.csv", reconciliation["checks"])

    readme = f"""# NeuroAI v2.2.0 → v2.3.0-dev migration baseline

Status: **{reconciliation["metadata"]["status"]}**

This package preserves and inventories the supplied v2.2.0 consolidated workbook,
consolidated report, and full programme archive. It creates deterministic sheet
exports and mechanical parity checks. It does not perform a substantive evidence
refresh or alter historical assessment decisions.

## Key outputs

- `MIGRATION_RECONCILIATION_REPORT_v2.3.0-rc1.json`
- `MIGRATION_RECONCILIATION_CHECKS_v2.3.0-rc1.csv`
- `workbook_sheet_inventory.json` / `.csv`
- `report_structure_inventory.json`
- `archive_inventory.json`
- `canonical_sheet_exports/`
- `MIGRATION_PACKAGE_MANIFEST.json`

## Gate

Static import parity allowed: `{reconciliation["next_phase_gate"]["static_import_parity_allowed"]}`

Substantive refresh allowed by this package alone: `False`

{BOUNDARY}
"""
    (args.output / "README.md").write_text(readme, encoding="utf-8")
    manifest = build_manifest(
        args.output,
        {
            "workbook": args.workbook,
            "report": args.report,
            "archive": args.archive,
            "sheet_map": args.sheet_map,
            "contract": args.contract,
        },
    )
    (args.output / "MIGRATION_PACKAGE_MANIFEST.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    print(json.dumps(reconciliation["metadata"], ensure_ascii=False, indent=2))
    return 0 if reconciliation["metadata"]["status"] == "PASS" else 1


if __name__ == "__main__":
    raise SystemExit(main())
