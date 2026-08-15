from __future__ import annotations

import io
import sys
import warnings
import zipfile
from pathlib import Path

import pytest
from docx import Document

from neuroai_workbench.products import archive as archive_module
from neuroai_workbench.products import docx as docx_module
from neuroai_workbench.products import pdf as pdf_module
from neuroai_workbench.products.query import query_release

REPO = Path(__file__).resolve().parents[2]
COMPACT = REPO / "examples" / "observatory" / "canonical_successor_snapshot_v1.7.json"


def _source_zip(
    entries: list[tuple[str, bytes]],
    *,
    date_time: tuple[int, int, int, int, int, int],
    archive_comment: bytes,
    member_comment: bytes,
    external_attr: int,
) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.comment = archive_comment
        for filename, payload in entries:
            info = zipfile.ZipInfo(filename=filename, date_time=date_time)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.create_system = 0
            info.external_attr = external_attr
            info.extra = b"\xfe\xca\x00\x00"
            info.comment = member_comment
            archive.writestr(info, payload)
    return buffer.getvalue()


def _assert_canonical_zip(payload: bytes) -> None:
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        infos = archive.infolist()
        assert archive.comment == b""
        assert [info.filename for info in infos] == sorted(info.filename for info in infos)
        assert infos
        for info in infos:
            assert info.date_time == archive_module.CANONICAL_ZIP_DATE
            assert info.compress_type == archive_module.CANONICAL_COMPRESSION
            assert info.create_system == 3
            assert info.create_version == 20
            assert info.extract_version == 20
            assert info.external_attr == archive_module.CANONICAL_EXTERNAL_ATTR
            assert info.extra == b""
            assert info.comment == b""


def test_shared_archive_canonicalization_erases_container_metadata() -> None:
    entries = [("b.txt", b"bravo"), ("a.txt", b"alpha")]
    first = _source_zip(
        entries,
        date_time=(2020, 1, 2, 3, 4, 6),
        archive_comment=b"first archive",
        member_comment=b"first member",
        external_attr=0o644 << 16,
    )
    second = _source_zip(
        list(reversed(entries)),
        date_time=(2026, 8, 15, 12, 34, 56),
        archive_comment=b"second archive",
        member_comment=b"second member",
        external_attr=0o755 << 16,
    )

    canonical_first = archive_module.canonicalize_zip_payload(first)
    canonical_second = archive_module.canonicalize_zip_payload(second)

    assert first != second
    assert canonical_first == canonical_second
    _assert_canonical_zip(canonical_first)


def test_shared_archive_rejects_duplicate_members() -> None:
    buffer = io.BytesIO()
    with warnings.catch_warnings():
        warnings.simplefilter("ignore", UserWarning)
        with zipfile.ZipFile(buffer, "w") as archive:
            archive.writestr("duplicate.txt", b"first")
            archive.writestr("duplicate.txt", b"second")

    with pytest.raises(ValueError, match="unique filenames"):
        archive_module.canonicalize_zip_payload(buffer.getvalue())


def test_shared_archive_can_preserve_explicit_member_order() -> None:
    payload = archive_module.render_deterministic_zip(
        [("b.txt", b"bravo"), ("a.txt", b"alpha")],
        sort_entries=False,
    )
    with zipfile.ZipFile(io.BytesIO(payload), "r") as archive:
        assert archive.namelist() == ["b.txt", "a.txt"]


def test_native_docx_is_byte_reproducible_and_canonical() -> None:
    query = query_release(COMPACT, depth="full", limit=None)
    first = docx_module.render_docx(query)
    second = docx_module.render_docx(query)

    assert first is not None
    assert second is not None
    assert first == second
    _assert_canonical_zip(first)

    document = Document(io.BytesIO(first))
    created = document.core_properties.created
    modified = document.core_properties.modified
    assert created is not None
    assert modified is not None
    assert created.replace(tzinfo=None) == archive_module.CANONICAL_DOCUMENT_TIME
    assert modified.replace(tzinfo=None) == archive_module.CANONICAL_DOCUMENT_TIME
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert query["release_sha256"] in text
    assert document.tables


def test_native_pdf_is_byte_reproducible_with_invariant_metadata() -> None:
    query = query_release(COMPACT, depth="full", limit=None)
    first = pdf_module.render_pdf(query)
    second = pdf_module.render_pdf(query)

    assert first is not None
    assert second is not None
    assert first == second
    assert first.startswith(b"%PDF")
    assert query["release_sha256"].encode("utf-8") in first
    assert b"/CreationDate (D:20000101000000" in first
    assert b"/ModDate (D:20000101000000" in first


def test_docx_stub_remains_deterministic_without_optional_dependency(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    query = query_release(COMPACT, depth="full", limit=None)
    monkeypatch.setitem(sys.modules, "docx", None)
    output = tmp_path / "report.docx"

    metadata = docx_module.write_docx(query, output)

    assert metadata["format"] == "docx-stub-text"
    assert query["release_sha256"] in output.read_text(encoding="utf-8")


def test_pdf_stub_remains_deterministic_without_optional_dependency(
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    query = query_release(COMPACT, depth="full", limit=None)
    monkeypatch.setattr(pdf_module, "render_pdf", lambda _query: None)
    output = tmp_path / "report.pdf"

    metadata = pdf_module.write_pdf(query, output)

    assert metadata["format"] == "pdf-stub-text"
    assert output.read_text(encoding="utf-8") == pdf_module.render_pdf_stub(query)
