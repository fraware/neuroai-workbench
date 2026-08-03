"""Native product-path assertions.

These tests require the optional product extras (openpyxl, python-docx, reportlab).
They must not soft-skip: CI installs hashed pins from requirements/constraints.txt so the
native XLSX/DOCX/PDF branches are exercised. Core/default installs remain free of these
extras for end users.
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import openpyxl  # noqa: F401 — required native extra; do not soft-skip
import reportlab  # noqa: F401 — required native extra; do not soft-skip
from docx import Document
from openpyxl import load_workbook

from neuroai_workbench.products.docx import write_docx
from neuroai_workbench.products.excel import render_analytical_workbook_bundle, write_analytical_workbook_bundle
from neuroai_workbench.products.generate import generate_publication_set, reconcile_formats
from neuroai_workbench.products.pdf import write_pdf
from neuroai_workbench.products.query import iter_appendix_sheets, query_release

REPO = Path(__file__).resolve().parents[2]
COMPACT = REPO / "examples" / "observatory" / "canonical_successor_snapshot_v1.7.json"
FULL = REPO / "examples" / "observatory" / "evidence_depth_release_v1.4.json"

# First-class full-depth sheets expected when present on the v1.4 fixture.
EXPECTED_FULL_SHEETS = {
    "release_summary",
    "coverage_counts",
    "organizations",
    "sources",
    "aliases",
    "organization_resolution",
    "regional_expansion",
    "ownership_capital_events",
    "models",
    "models_datasets",
    "trial_sites",
    "participant_authority",
    "suppliers",
    "data_quality_findings",
    "coverage_exit_conditions",
    "methodology_source_universes",
    "projection_limits",
    "verification",
}

EXPECTED_COMPACT_SHEETS = {
    "release_summary",
    "successor_counts",
    "baseline_counts",
    "reopening_decisions",
    "delta_records",
    "delta_counts",
    "assessment_successor_delta",
    "provenance_links",
    "verification",
}


def test_native_workbook_has_single_verification_sheet() -> None:
    query = query_release(COMPACT, depth="full", limit=None)
    payload = render_analytical_workbook_bundle(query)
    workbook = load_workbook(io.BytesIO(payload))
    assert "verification" in workbook.sheetnames
    assert "verification1" not in workbook.sheetnames
    verification_like = [name for name in workbook.sheetnames if name.startswith("verification")]
    assert verification_like == ["verification"]
    flat = [item for row in workbook["verification"].iter_rows(values_only=True) for item in row]
    assert query["release_sha256"] in flat
    assert "openpyxl-native-xlsx" in flat


def test_native_workbook_write_reports_openpyxl_format(tmp_path: Path) -> None:
    query = query_release(COMPACT, depth="full", limit=None)
    output = tmp_path / "analytical-workbook.xlsx"
    meta = write_analytical_workbook_bundle(query, output)
    assert meta["format"] == "openpyxl-native-xlsx"
    workbook = load_workbook(output)
    assert workbook.sheetnames.count("verification") == 1
    assert "verification1" not in workbook.sheetnames


def test_native_full_release_workbook_sheet_coverage() -> None:
    query = query_release(FULL, depth="full", limit=None)
    assert EXPECTED_FULL_SHEETS <= set(query["rows"])
    # Absent canonical sections must not be invented.
    for absent in ("systems", "captures", "candidates", "adjudications", "evidence_register"):
        assert absent not in query["rows"]
    assert query["rows"]["aliases"]
    assert {"organization_id", "alias"} <= set(query["rows"]["aliases"][0])
    assert query["rows"]["models"]
    assert query["rows"]["models_datasets"]

    payload = render_analytical_workbook_bundle(query)
    workbook = load_workbook(io.BytesIO(payload))
    assert "verification1" not in workbook.sheetnames
    assert workbook.sheetnames.count("verification") == 1
    # Excel truncates titles to 31 characters.
    projected = {name[:31] for name in query["rows"]}
    assert projected <= set(workbook.sheetnames)


def test_native_compact_successor_workbook_sheet_coverage() -> None:
    query = query_release(COMPACT, depth="full", limit=None)
    assert EXPECTED_COMPACT_SHEETS <= set(query["rows"])
    assert query["rows"]["delta_records"]
    assert query["rows"]["provenance_links"]
    payload = render_analytical_workbook_bundle(query)
    workbook = load_workbook(io.BytesIO(payload))
    assert workbook.sheetnames.count("verification") == 1
    assert "verification1" not in workbook.sheetnames
    projected = {name[:31] for name in query["rows"]}
    assert projected <= set(workbook.sheetnames)


def test_native_docx_is_real_office_container(tmp_path: Path) -> None:
    query = query_release(COMPACT, depth="full", limit=None)
    output = tmp_path / "current-state-report.docx"
    meta = write_docx(query, output)
    assert meta["format"] == "python-docx-native"
    payload = output.read_bytes()
    assert payload[:2] == b"PK"
    with zipfile.ZipFile(io.BytesIO(payload)) as archive:
        names = set(archive.namelist())
    assert "[Content_Types].xml" in names
    assert "word/document.xml" in names
    document = Document(str(output))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert query["release_sha256"] in text
    assert "Withheld claims" in text
    assert "Generators are views" in text
    assert document.tables, "DOCX must render substantive appendix tables from the query graph"
    appendix_names = {name for name, _ in iter_appendix_sheets(query)}
    assert appendix_names
    assert any("delta" in paragraph.text.lower() for paragraph in document.paragraphs)


def test_native_docx_full_release_has_substantive_appendices(tmp_path: Path) -> None:
    query = query_release(FULL, depth="full", limit=None)
    output = tmp_path / "full-report.docx"
    meta = write_docx(query, output)
    assert meta["format"] == "python-docx-native"
    document = Document(str(output))
    assert len(document.tables) >= 5
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "organizations" in text.lower() or "Publication appendices" in text
    assert "No regulatory authorization" in text


def test_native_pdf_is_real_pdf_container(tmp_path: Path) -> None:
    query = query_release(COMPACT, depth="full", limit=None)
    output = tmp_path / "current-state-report.pdf"
    meta = write_pdf(query, output)
    assert meta["format"] == "reportlab-native-pdf"
    payload = output.read_bytes()
    assert payload.startswith(b"%PDF")
    assert query["release_sha256"].encode("utf-8") in payload
    assert b"Withheld claims" in payload or b"withheld" in payload.lower()
    assert b"Generators are views" in payload or b"generators are views" in payload.lower()
    # Substantive multi-page appendix projection (not a one-page extract).
    assert payload.count(b"/Type /Page") >= 2 or payload.count(b"/Type/Page") >= 2


def test_native_pdf_full_release_is_multi_page(tmp_path: Path) -> None:
    query = query_release(FULL, depth="full", limit=None)
    output = tmp_path / "full-report.pdf"
    meta = write_pdf(query, output)
    assert meta["format"] == "reportlab-native-pdf"
    payload = output.read_bytes()
    assert payload.startswith(b"%PDF")
    assert query["release_sha256"].encode("utf-8") in payload
    assert payload.count(b"/Type /Page") >= 3 or payload.count(b"/Type/Page") >= 3


def test_generate_publication_set_uses_native_formats(tmp_path: Path) -> None:
    report = generate_publication_set(COMPACT, tmp_path, limit=None, depth="full")
    products = report["products"]
    assert products["analytical_workbook"]["format"] == "openpyxl-native-xlsx"
    assert products["docx"]["format"] == "python-docx-native"
    assert products["pdf"]["format"] == "reportlab-native-pdf"

    workbook = load_workbook(tmp_path / "analytical-workbook.xlsx")
    assert "verification" in workbook.sheetnames
    assert "verification1" not in workbook.sheetnames
    assert [name for name in workbook.sheetnames if name.startswith("verification")] == ["verification"]

    docx_bytes = (tmp_path / "current-state-report.docx").read_bytes()
    assert docx_bytes[:2] == b"PK"
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as archive:
        assert "word/document.xml" in archive.namelist()

    pdf_bytes = (tmp_path / "current-state-report.pdf").read_bytes()
    assert pdf_bytes.startswith(b"%PDF")

    query = query_release(COMPACT, depth="full", limit=None)
    reconciliation = reconcile_formats(query, products)
    assert reconciliation["reconciled"] is True


def test_generate_publication_set_full_release_native(tmp_path: Path) -> None:
    report = generate_publication_set(FULL, tmp_path / "full", limit=None, depth="full")
    products = report["products"]
    assert products["analytical_workbook"]["format"] == "openpyxl-native-xlsx"
    assert products["docx"]["format"] == "python-docx-native"
    assert products["pdf"]["format"] == "reportlab-native-pdf"
    workbook = load_workbook(tmp_path / "full" / "analytical-workbook.xlsx")
    assert workbook.sheetnames.count("verification") == 1
    assert "aliases" in workbook.sheetnames or "aliases"[:31] in workbook.sheetnames
    query = query_release(FULL, depth="full", limit=None)
    assert reconcile_formats(query, products)["reconciled"] is True
